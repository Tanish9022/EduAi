import os
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from typing import List
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from PIL import Image
import google.generativeai as genai
from app.core.config import settings

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 100):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False,
        )
        if settings.GEMINI_API_KEY:
            genai.configure(api_key=settings.GEMINI_API_KEY)

    def process_file(self, file_path: str) -> List[Document]:
        """Unified method to process .txt, .pdf, .docx, and image files."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".txt":
            return self.process_text_file(file_path)
        elif ext == ".pdf":
            return self.process_pdf_file(file_path)
        elif ext == ".docx":
            return self.process_docx_file(file_path)
        elif ext in [".png", ".jpg", ".jpeg"]:
            return self.process_image_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def process_text_file(self, file_path: str) -> List[Document]:
        """Loads a text file and splits it into chunks."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            
        metadata = {"source": os.path.basename(file_path)}
        docs = [Document(page_content=text, metadata=metadata)]
        return self.text_splitter.split_documents(docs)

    def process_pdf_file(self, file_path: str) -> List[Document]:
        """Loads a PDF file and splits it into chunks. If a page is scanned, uses Gemini to read it."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        import io
        doc = fitz.open(file_path)
        text = ""
        for i, page in enumerate(doc):
            page_text = page.get_text()
            if len(page_text.strip()) < 50:
                if settings.GEMINI_API_KEY:
                    try:
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        with Image.open(io.BytesIO(img_data)) as img:
                            img.load()
                            model = genai.GenerativeModel('gemini-3.5-flash')
                            prompt = (
                                "This is a page from a scanned PDF document. Extract all readable text, tables, "
                                "dates, and structure. Keep the output clean, accurate, and in the original language."
                            )
                            response = model.generate_content([prompt, img])
                            page_text = response.text
                    except Exception as e:
                        print(f"Warning: Gemini failed to transcribe scanned page {i+1}: {e}")
                else:
                    print("Warning: GEMINI_API_KEY not configured, cannot OCR scanned PDF page.")
            
            text += f"\n--- Page {i+1} ---\n" + page_text
            
        metadata = {"source": os.path.basename(file_path)}
        docs = [Document(page_content=text, metadata=metadata)]
        return self.text_splitter.split_documents(docs)

    def process_docx_file(self, file_path: str) -> List[Document]:
        """Loads a DOCX file and splits it into chunks."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
            
        metadata = {"source": os.path.basename(file_path)}
        docs = [Document(page_content=text, metadata=metadata)]
        return self.text_splitter.split_documents(docs)

    def process_image_file(self, file_path: str) -> List[Document]:
        """Loads an image file, uses Gemini Vision to describe/extract text, and splits it into chunks."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        if not settings.GEMINI_API_KEY:
            raise ValueError("GEMINI_API_KEY not configured. Cannot process image files.")

        try:
            with Image.open(file_path) as img:
                img.load()
                model = genai.GenerativeModel('gemini-2.5-flash')
                
                prompt = (
                    "Describe this document image in full detail. Extract all readable text, tables, graphs, "
                    "forms, dates, structured data, and general context. Keep the format clean and highly detailed, "
                    "so that it can be searched and referenced by an AI admissions assistant."
                )
                
                response = model.generate_content([prompt, img])
                text = response.text
        except Exception as e:
            raise RuntimeError(f"Failed to process image with Gemini: {str(e)}")

        metadata = {"source": os.path.basename(file_path)}
        docs = [Document(page_content=text, metadata=metadata)]
        return self.text_splitter.split_documents(docs)
